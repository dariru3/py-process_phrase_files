import re

from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from .config_loader import CONFIG


def format_subscripts(paragraph):
    """Helper function to format text with subscript tags"""
    text = paragraph.text
    # Split on subscript tags: {_{>number<}_{}} as example, adapt pattern to whatever subscript tags are
    # Using {_{>...<}_} style would be similar to superscript but with subscript tags, example assumed {v>2<v} or similar
    # Your current tags to remove are {_> number <_}, so here we look for those to format subscript
    parts = re.split(r"(\{_>.*?<_\})", text)

    segments = []
    for part in parts:
        if part.startswith("{_>") and part.endswith("<_}"):
            # This is a subscript tag
            subscript_text = part[3:-3]  # Remove the tags to get the number/text inside
            segments.append((subscript_text, True))
        else:
            segments.append((part, False))

    # Remove original runs once before rebuilding.
    for run in list(paragraph.runs):
        p = run._element.getparent()
        p.remove(run._element)

    # Append new runs with formatting.
    for text_part, is_subscript in segments:
        run = paragraph.add_run(text_part)
        if is_subscript:
            run.font.subscript = True


def format_superscripts(paragraph):
    """Helper function to format text with superscript tags"""

    def apply_superscript(run, text):
        """Helper function to set text as superscript"""
        run.text = text
        run.font.superscript = True

    text = paragraph.text
    parts = re.split(r"(\{\^\>.*?\<\^\}|\{.*?\>.*?\<.*?\})", text)

    new_runs = []
    for part in parts:
        if part.startswith("{^>") and part.endswith("<^}"):
            # This is a superscript
            superscript_text = part[3:-3]
            run = paragraph.add_run()
            apply_superscript(run, superscript_text)
            new_runs.append(run)
        elif re.match(r"\{.*?\>.*?\<.*?\}", part):
            # Handle other custom tags if necessary
            inner_text = part.split(">")[1].split("<")[0]
            run = paragraph.add_run()
            apply_superscript(run, inner_text)
            new_runs.append(run)
        else:
            # Normal text
            run = paragraph.add_run(part)
            new_runs.append(run)

    # Remove original runs
    for run in paragraph.runs:
        p = run._element.getparent()
        p.remove(run._element)

    # Append new runs with formatting
    for new_run in new_runs:
        run = paragraph.add_run(new_run.text)
        run.font.superscript = new_run.font.superscript


def reformat_text(table):
    """Process a single table in the document"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # format_superscripts(paragraph)
                format_subscripts(paragraph)


def change_cell_color(cells, background_color):
    for cell in cells:
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), background_color)
        tcPr.append(shd)


def set_column_language(table, column_index, language_code):
    """
    Not working as intended
    """
    for row in table.rows:
        cell = row.cells[column_index]
        for paragraph in cell.paragraphs:
            rPr = paragraph.runs[0].element.get_or_add_rPr()
            lang = OxmlElement("w:lang")
            lang.set(qn("w:val"), language_code)
            rPr.append(lang)


# format column widths and header row cell color
def format_table(table):
    t_settings = CONFIG["TableFormattingSettings"]
    table.style = "Table Grid"
    row_widths = t_settings["RowWidths"]

    for i, width in enumerate(row_widths):
        for cell in table.columns[i].cells:
            cell.width = Mm(width)

    blue_color = "95B3D7"  # Hex code for blue
    first_column_cells = table.rows[0].cells
    change_cell_color(first_column_cells, blue_color)


def apply_conditional_formatting(table):
    """
    Change row cell colors to gray if either condition is met.
    Condition 1: There is text in the target cell and the match is either "100" or "101"
    Condition 2: There is text in the target cell and the comment is "lock" or "locked"
    """
    c_settings = CONFIG["ConditionalFormattingSettings"]
    target_column_index = c_settings["TargetColumnIndex"]  # 2
    match_column_index = c_settings["MatchColumnIndex"]  # 3
    comment_column_index = c_settings["CommentColumnIndex"]  # 4
    comment_to_gray = c_settings["CommentToGray"]  # ['lock', 'locked']
    match_to_gray = c_settings["MatchToGray"]  # ['100', '101']
    background_color = c_settings["BackgroundColor"]  # "D9D9D9"

    for row in table.rows:
        target_value = row.cells[target_column_index].text.strip()
        match_value = row.cells[match_column_index].text.strip()
        comment_value = row.cells[comment_column_index].text.lower().strip()

        condition_1_met = target_value and match_value in match_to_gray
        condition_2_met = target_value and comment_value in comment_to_gray

        if condition_1_met or condition_2_met:
            cells_to_color = [cell for cell in row.cells]
            change_cell_color(cells_to_color, background_color)


def set_landscape_orientation(document):
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)  # A4 width
    section.page_height = Mm(210)  # A4 height


# set font size and line spacing
def format_font_lines(document):
    style = document.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    line_space = 1.15
    column_1_font_size = 8

    for paragraph in document.paragraphs:
        apply_paragraph_format(paragraph, style, line_space)

    for table in document.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if i == 0:  # set font size of column 1 to 8pt
                        apply_paragraph_format(
                            paragraph, style, line_space, column_1_font_size
                        )
                    else:
                        apply_paragraph_format(paragraph, style, line_space)


# helper function for format_font_lines()
def apply_paragraph_format(paragraph, style, line_space, font_size=None):
    paragraph.style = style
    paragraph.paragraph_format.line_spacing = line_space
    if font_size:
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(font_size)


def apply_formatting_pipe(table, doc):
    format_table(table)
    apply_conditional_formatting(table)
    reformat_text(table)  # apply subscript
    set_column_language(table, 1, "ja-JP")
    set_landscape_orientation(doc)
    format_font_lines(doc)
