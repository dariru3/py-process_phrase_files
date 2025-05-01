from docx import Document
from .table_to_df import table_to_df
from .save_formatting import extract_formatting_from_column
from .process_mxliff import remove_tags
from .config_loader import CONFIG

def delete_first_n_tables(doc, n):
    for _ in range(n):
        if len(doc.tables) > 0:
            table = doc.tables[0]
            table._element.getparent().remove(table._element)

def copy_content_to_table(original_table, new_table, columns_to_copy):
    for row in original_table.rows:
        new_row = new_table.add_row()
        new_cells = new_row.cells
        for i, col_index in enumerate(columns_to_copy):
            original_text = row.cells[col_index].text
            cleansed_text = remove_tags(original_text)
            new_cells[i].text = cleansed_text

def process_word_file(file_path, output_folder):
    # Load settings
    p_settings = CONFIG["ProcessingDocSettings"]
    final_col_length = len(CONFIG["GeneralSettings"]["Column_Headers"])

    print("Processing .DOCX file...")
    doc = Document(file_path)

    # Remove other tables
    tables_to_delete = p_settings["DeleteFirstNTables"]
    delete_first_n_tables(doc, tables_to_delete)

    # Get contents from source table from certain columns
    original_table = doc.tables[0]
    new_table = doc.add_table(rows=0, cols=final_col_length)
    columns_to_copy = p_settings["ColumnsToKeep"]
    copy_content_to_table(original_table, new_table, columns_to_copy)

    # Save text formatting to reapply later
    formatting_info = extract_formatting_from_column(doc, 1, [1, 2])

    print("Preview of new_table via python-docx:")
    for i, row in enumerate(new_table.rows[:5]):
        cell_texts = [cell.text for cell in row.cells]
        print(f" Row {i}: {cell_texts}")

    df_table = table_to_df(new_table)

    return df_table, formatting_info
