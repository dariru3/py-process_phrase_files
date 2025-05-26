from docx.shared import Pt, RGBColor
from src.process_mxliff import remove_tags

def extract_formatting_from_column(doc, table_num, col_nums):
    table = doc.tables[table_num]
    formatting_info = {}

    for row_idx, row in enumerate(table.rows):
        row_format = {}
        for col in col_nums:
            cell = row.cells[col]
            cell_info = []
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    clean_text = remove_tags(run.text)
                    run_info = {
                        "text": clean_text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name,
                        "font_size": run.font.size.pt if run.font.size else None,
                        "font_color": run.font.color.rgb if run.font.color else None,
                        "superscript": run.font.superscript,
                        "subscript": run.font.subscript,
                    }
                    cell_info.append(run_info)
            row_format[col] = cell_info
        formatting_info[row_idx] = row_format

    return formatting_info

def reapply_formatting_to_column(table, formatting_info, col_nums, table_num=0):
    """Reapply saved formatting to the target table.

    Parameters
    ----------
    table : docx.table.Table
        Destination table where formatting should be restored.
    formatting_info : dict
        Mapping of row and column indices to formatting attributes as
        returned by ``extract_formatting_from_column``.
    col_nums : Iterable[int]
        Original column indices that ``formatting_info`` was extracted from.
        The order of ``col_nums`` determines the target columns in ``table``
        starting from column ``1`` (column ``0`` is reserved for the index
        column created during processing).
    table_num : int, optional
        Currently unused but kept for backward compatibility.
    """

    col_mapping = {orig_col: idx + 1 for idx, orig_col in enumerate(col_nums)} # { 3: 1, 5: 2 }
    for row_idx, cols_info in formatting_info.items():
        for orig_col, new_col in col_mapping.items():
            cell_info = cols_info.get(orig_col, [])
            has_previous_text = any(run_info["text"] for run_info in cell_info)
            if not has_previous_text:
                continue

            cell = table.cell(row_idx + 1, new_col) # Skip header row
            cell.text = ""
            for run_info in cell_info:
                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                run = paragraph.add_run(run_info["text"])
                run.bold = run_info.get("bold")
                run.italic = run_info.get("italic")
                run.underline = run_info.get("underline")
                run.font.name = run_info.get("font_name")
                if run_info.get("font_size"):
                    run.font.size = Pt(run_info["font_size"])
                if run_info.get("font_color"):
                    run.font.color.rgb = RGBColor.from_string(str(run_info["font_color"]))
                run.font.superscript = run_info.get("superscript")
                run.font.subscript = run_info.get("subscript")
