
# Start of config_loader.py
import json

def load_config(config_path = 'config.json'):
    with open(config_path, 'r') as config_file:
        return json.load(config_file)

CONFIG = load_config()
# End of config_loader.py


# Start of process_mxliff.py
import xml.etree.ElementTree as ET
import pandas as pd

def parse_mxliff_to_df(mxliff_file):
    print("Processing .MXLIFF file...")
    # Register the namespace to properly handle prefixed attributes
    ET.register_namespace('m', 'urn:oasis:names:tc:xliff:document:1.2')

    # Parse the MXLIFF file
    tree = ET.parse(mxliff_file)
    root = tree.getroot()

    # Define the namespaces used in your MXLIFF file
    namespaces = {'m': 'urn:oasis:names:tc:xliff:document:1.2'}

    # Initialize lists to hold the extracted data
    sources = []
    targets = []
    match_qualities = []

    # Loop through each translation unit in the MXLIFF file
    for trans_unit in root.findall('.//m:trans-unit', namespaces):
        source_text = trans_unit.find('m:source', namespaces).text if trans_unit.find('m:source', namespaces) is not None else ''
        target_text = trans_unit.find('m:target', namespaces).text if trans_unit.find('m:target', namespaces) is not None else ''
        match_quality = '0' # Default value
        
        # Check for alt-trans elements with origin="memsource-tm" and extract match-quality
        for alt_trans in trans_unit.findall('.//m:alt-trans', namespaces):
            if alt_trans.attrib.get('origin') == 'memsource-tm':
                match_quality = alt_trans.attrib.get('match-quality', '0')
                match_quality = int(float(match_quality) * 100)
                break  # Assuming we only need the first matching alt-trans entry

        sources.append(source_text)
        targets.append(target_text)
        match_qualities.append(match_quality)

    # Create a DataFrame from the extracted data
    df = pd.DataFrame({
        'Source': sources,
        'Target': targets,
        'Match': match_qualities
    })

    df.index += 1
    df.reset_index(inplace=True)
    df.rename(columns={'index': 'Index'}, inplace=True)

    df['Comment'] = ""

    return df
# End of process_mxliff.py


# Start of format_helper.py
from docx.shared import Mm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from config_loader import CONFIG

def change_cell_color(cells, background_color):
    for cell in cells:
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), background_color)
        tcPr.append(shd)

def set_column_language(table, column_index, language_code):
    for row in table.rows:
        cell = row.cells[column_index]
        for paragraph in cell.paragraphs:
            rPr = paragraph.runs[0].element.get_or_add_rPr()
            lang = OxmlElement('w:lang')
            lang.set(qn('w:val'), language_code)
            rPr.append(lang)

# format column widths and header row cell color
def format_table(table):
    t_settings = CONFIG["TableFormattingSettings"]
    table.style = 'Table Grid'
    row_widths = t_settings["RowWidths"]

    for i, width in enumerate(row_widths):
        for cell in table.columns[i].cells:
            cell.width = Mm(width)

    blue_color = "95B3D7"  # Hex code for blue
    first_column_cells = table.rows[0].cells
    change_cell_color(first_column_cells, blue_color)

def apply_conditional_formatting(table):
    '''
    Change row cell colors to gray if either condition is met.
    Condition 1: There is text in the target cell and the match is either "100" or "101"
    Condition 2: There is text in the target cell and the comment is "lock" or "locked"
    '''
    c_settings = CONFIG["ConditionalFormattingSettings"]
    target_column_index = c_settings["TargetColumnIndex"] # 2
    match_column_index = c_settings["MatchColumnIndex"] #3
    comment_column_index = c_settings["CommentColumnIndex"] # 4
    comment_to_gray = c_settings["CommentToGray"] # ['lock', 'locked']
    match_to_gray = c_settings["MatchToGray"] # ['100', '101']
    background_color= c_settings["BackgroundColor"] # "D9D9D9"

    for row in table.rows:
        target_value = row.cells[target_column_index].text.strip()
        match_value = row.cells[match_column_index].text.strip()
        comment_value = row.cells[comment_column_index].text.lower().strip()

        condition_1_met = target_value and match_value in match_to_gray
        condition_2_met = target_value and comment_value in comment_to_gray

        if condition_1_met or condition_2_met:
            cells_to_color = [cell for cell in row.cells]
            change_cell_color(cells_to_color, background_color)

def set_landscape_orientation(document):
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297) # A4 width
    section.page_height = Mm(210) # A4 height

# set font size and line spacing
def format_font_lines(document):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    line_space = 1.15
    column_1_font_size = 8

    for paragraph in document.paragraphs:
        apply_paragraph_format(paragraph, style, line_space)

    for table in document.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if i == 0: # set font size of column 1 to 8pt
                        apply_paragraph_format(paragraph, style, line_space, column_1_font_size)
                    else:
                        apply_paragraph_format(paragraph, style, line_space)

# helper function for format_font_lines()
def apply_paragraph_format(paragraph, style, line_space, font_size=None):
    paragraph.style = style
    paragraph.paragraph_format.line_spacing = line_space
    if font_size:
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(font_size)
# End of format_helper.py


# Start of process_word.py
from docx import Document
from table_to_df import table_to_df
import os
import re
from config_loader import CONFIG

def delete_first_n_tables(doc, n):
    for _ in range(n):
        if len(doc.tables) > 0:
            table = doc.tables[0]
            table._element.getparent().remove(table._element)

def copy_content_to_table(original_table, new_table, columns_to_copy):
    for row in original_table.rows:
        new_row = new_table.add_row()
        new_cells = new_row.cells
        for i, col_index in enumerate(columns_to_copy):
            new_cells[i].text = row.cells[col_index].text

def process_word_file(file_path, output_folder, attempts=1):
    p_settings = CONFIG["ProcessingSettings"]
    final_col_length = len(CONFIG["GeneralSettings"]["Column_Headers"])
    if attempts == 1:
        print("Processing .DOCX file...")
    max_attempts = p_settings["MaxAttempts"]
    doc = Document(file_path)

    tables_to_delete = p_settings["DeleteFirstNTables"]
    delete_first_n_tables(doc=doc, n=tables_to_delete)

    columns_to_copy = adjust_columns_by_attempts(attempts, p_settings)

    original_table = doc.tables[0]
    new_table = doc.add_table(rows=0, cols=final_col_length)

    copy_content_to_table(original_table, new_table, columns_to_copy)

    if validate_table_contents(new_table, p_settings):
        df_table = table_to_df(new_table)
        print("Success!")
        return df_table
    else:
        if attempts < max_attempts:
            print(f'Attempt {attempts} failed, trying again...')
            return process_word_file(file_path, output_folder, attempts + 1)
        else:
            print(f'Maximum attempts reached for file {file_path}. File processing aborted.')
            return None

def contains_japanese(text, process_settings):
    # Regular expression for matching Japanese characters
    pattern = process_settings["JapanesePattern"]
    return re.search(pattern, text) is not None

def validate_table_contents(new_table, process_settings):
    valid_rows = True
    for i, row in enumerate(new_table.rows[1:11]):
        column_3_target_text = row.cells[2].text

        if column_3_target_text and contains_japanese(column_3_target_text, process_settings):
            print(f"Invalid row {i}: {column_3_target_text}")
            valid_rows = False
    
    return valid_rows

def adjust_columns_by_attempts(attempts, process_settings):
    attempt_1_col = process_settings["Mapping_1"]
    attempt_2_col = process_settings["Mapping_2"]
    attempts_mapping = {
        1: ("First attempt", attempt_1_col),
        2: ("Second attempt", attempt_2_col),
    }

    message, columns = attempts_mapping.get(attempts, ("Second attempt failed", None))
    print(message)
    return columns
# End of process_word.py


# Start of df_to_word.py
from docx import Document
import os
import pandas as pd
from process_mxliff import parse_mxliff_to_df
import format_helper as help
from merge_df import merge_dfs
from process_word import process_word_file

def get_file_pairs(folder_path):
    docx_files = {}
    mxliff_files = {}
    for filename in os.listdir(folder_path):
        base_name, ext = os.path.splitext(filename)
        if ext == ".docx":
            docx_files[base_name] = filename
        elif ext == ".mxliff":
            mxliff_files[base_name] = filename
    
    # Pairing files with the same base name
    pairs = []
    for base_name, docx_file in docx_files.items():
        mxliff_file = mxliff_files.get(base_name)
        if mxliff_file:
            pairs.append((docx_file, mxliff_file))
    return pairs

def dataframe_to_word_table(docx_file, df, output_folder):
    doc = Document()
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.autofit = False

    # Rename column headers
    df.rename(columns={'Index': 'p', 'Source': 'Japanese', 'Target': 'English'}, inplace=True)

    # Add header row
    for i, column in enumerate(df.columns):
        table.cell(0, i).text = str(column)
    
    # Add data rows
    for index, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if pd.isnull(value) or value == "None":
                cells[i].text = ""
            else:
                cells[i].text = str(value)
    
    help.format_table(table)
    help.apply_conditional_formatting(table)
    help.set_landscape_orientation(doc)
    help.format_font_lines(doc)
    help.set_column_language(table, 1, 'ja-JP')
    
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    output_file_path = os.path.join(output_folder, f"{os.path.splitext(docx_file)[0]}_merged.docx")
    doc.save(output_file_path)
    print(f"Merged tables saved as Word document: {output_file_path}.")

def process_files(docx_file, mxliff_file, input_folder, output_folder):
    # Process the Word and MXLIFF files
    df_word = process_word_file(os.path.join(input_folder, docx_file), output_folder)
    if not df_word.empty:
        df_mxliff = parse_mxliff_to_df(os.path.join(input_folder, mxliff_file))
    else:
        print("Failed to process Word file.")
        return

    # Merge the DataFrames
    merged_df = merge_dfs(df_word, df_mxliff)

    # Save the merged DataFrame to a Word document
    dataframe_to_word_table(docx_file, merged_df, output_folder)
# End of df_to_word.py


# Start of merge_df.py
import pandas as pd

def merge_dfs(df1, df2):
    # Convert Index to int
    df1['Index'] = df1['Index'].astype(int)
    df2['Index'] = df2['Index'].astype(int)
    # Convert Match to int, with error handling
    df1['Match'] = pd.to_numeric(df1['Match'], errors='coerce').fillna(0).astype(int)
    df2['Match'] = pd.to_numeric(df2['Match'], errors='coerce').fillna(0).astype(int)

    df_combined = pd.merge(df1, df2, on=['Index', 'Source'], how='outer', suffixes=('', '_df2'))

    # Now, select the best values for each column based on availability and preference
    df_combined['Target'] = df_combined['Target'].where(df_combined['Target'] != '', df_combined['Target_df2'])
    df_combined['Match'] = df_combined['Match'].fillna(0).astype(int)
    df_combined['Match'] = df_combined['Match'].where(df_combined['Match'] != 0, df_combined['Match_df2']).fillna(0).astype(int)
    df_combined['Comment'] = df_combined['Comment'].where(df_combined['Comment'] != '', df_combined['Comment_df2'])

    # Drop the temporary columns from df2
    df_combined.drop(columns=['Target_df2', 'Match_df2', 'Comment_df2'], inplace=True)

    return df_combined
# End of merge_df.py


# Start of main.py
import os
from config_loader import CONFIG
from df_to_word import get_file_pairs, process_files

def filter_unprocessed_pairs(pairs, output_folder):
    unprocessed_pairs = []
    for docx_file, mxliff_file in pairs:
        # Extract the base name from the docx file (base name for mxliff file should be the same)
        base_name, _ = os.path.splitext(os.path.basename(docx_file))

        # Check if the merged file exists
        merged_filename = f"{base_name}_merged.docx"
        merged_file_path = os.path.join(output_folder, merged_filename)
        if not os.path.exists(merged_file_path):
            unprocessed_pairs.append((docx_file, mxliff_file))
        else:
            print(f"Skipped processing for {base_name} because the merged file already exists.")
    return unprocessed_pairs

def main():
    g_settings = CONFIG["GeneralSettings"]
    input_folder = g_settings["InputFolderPath"] # "input_files/"
    output_folder = g_settings["OutputFolderPath"] # "output_files/"

    pairs = get_file_pairs(input_folder)
    unprocessed_pairs = filter_unprocessed_pairs(pairs, output_folder)
    for docx_file, mxliff_file in unprocessed_pairs:
        print(f"File pair:\n{docx_file}\n{mxliff_file}")
        process_files(docx_file, mxliff_file,input_folder, output_folder)

if __name__ == "__main__":
    main()
# End of main.py


# Start of table_to_df.py
import pandas as pd
from config_loader import CONFIG

def table_to_df(table):
    column_headers = CONFIG["GeneralSettings"]["Column_Headers"]
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        data.append(row_data)
    return pd.DataFrame(data, columns=column_headers)
# End of table_to_df.py

