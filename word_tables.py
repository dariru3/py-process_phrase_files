from docx import Document
import os

def process_word_file(file_path):
    doc = Document(file_path)

    # Delete the first three tables
    for _ in range(3):
        if len(doc.tables) > 0:
            tbl = doc.tables[0]
            tbl._element.getparent().remove(tbl._element)

    # Process the fourth table (now the first table in the document)
    if len(doc.tables) > 0:
        original_table = doc.tables[0]
        rows = original_table.rows

        # Determine the number of columns to be copied
        num_cols = len(rows[0].cells) - 5  # Subtract first 3 and last 2 columns

        # Create a new table with the desired number of columns
        new_table = doc.add_table(rows=0, cols=num_cols)

        for row in rows:
            new_row = new_table.add_row()
            new_cells = new_row.cells
            # Copy the content from the middle columns of the original table
            for idx, cell in enumerate(row.cells[3:-2]):
                new_cells[idx].text = cell.text

        # Remove the original table
        original_table._element.getparent().remove(original_table._element)

    # Save the document
    new_file_path = file_path.replace('.docx', '_processed.docx')
    doc.save(new_file_path)

    print(f"Processed file saved as: {new_file_path}")

def process_all_word_files_in_folder(folder_path):
    for file_name in os.listdir(folder_path):
        if file_name.endswith('.docx'):
            file_path = os.path.join(folder_path, file_name)
            process_word_file(file_path)

# Example usage
input_folder = 'input_files'
process_all_word_files_in_folder(input_folder)
