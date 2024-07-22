from docx import Document
from docx.shared import Pt, RGBColor

def extract_formatting_from_column(file_path, table_num, col_num):
    doc = Document(file_path)
    table = doc.tables[table_num]
    formatting_info = {}

    for row_idx, row in enumerate(table.rows):
        cell = row.cells[col_num]
        cell_info = []
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run_info = {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": run.font.size.pt if run.font.size else None,
                    "font_color": run.font.color.rgb if run.font.color else None,
                    "superscript": run.font.superscript,
                    "subscript": run.font.subscript,
                }
                cell_info.append(run_info)
        formatting_info[row_idx] = cell_info

    return formatting_info

# Example usage
file_path = 'input_files/企業価値向上（新規依頼）-ja-en-T.docx'
table_num = 3  # 3rd table (index starts at 0)
col_num = 5    # 5th column
formatting_info = extract_formatting_from_column(file_path, table_num, col_num)
print(formatting_info)

def reapply_formatting_to_column(file_path, table_num, col_num, formatting_info):
    doc = Document(file_path)
    table = doc.tables[table_num]

    for row_idx, cell_info in formatting_info.items():
        cell = table.cell(row_idx + 1, col_num) # +1 = start 2nd row
        cell.text = ""
        for run_info in cell_info:
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            run = paragraph.add_run(run_info["text"])
            run.bold = run_info.get("bold")
            run.italic = run_info.get("italic")
            run.underline = run_info.get("underline")
            run.font.name = run_info.get("font_name")
            if run_info.get("font_size"):
                run.font.size = Pt(run_info["font_size"])
            if run_info.get("font_color"):
                run.font.color.rgb = RGBColor.from_string(run_info["font_color"])
            run.font.superscript = run_info.get("superscript")
            run.font.subscript = run_info.get("subscript")

    doc.save('output_files/output_formatted.docx')

# Example usage
file_path = 'output_files/企業価値向上（新規依頼）-ja-en-T_merged.docx'
table_num = 0  # 1st table (index starts at 0)
col_num = 2    # 3rd column
reapply_formatting_to_column(file_path, table_num, col_num, formatting_info)
