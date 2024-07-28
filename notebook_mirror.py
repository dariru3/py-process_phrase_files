import os
import json
import xml.etree.ElementTree as ET
import pandas as pd
from docx import Document
from docx.shared import Mm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

# Configuration loading function
def load_config(config_path='/content/drive/MyDrive/MagicBox/configFiles/config.json'):
    with open(config_path, 'r') as config_file:
        return json.load(config_file)

CONFIG = load_config()

# Function to ensure a folder exists
def ensure_folder_exists(folder_path):
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        print(f"Created folder: {folder_path}")

# Function to check for required files
def check_for_files(directory_path):
    has_docx = False
    has_mxliff = False
    for item in os.listdir(directory_path):
        item_path = os.path.join(directory_path, item)
        if os.path.isfile(item_path):  # Ensure the item is a file
            if item.endswith('.docx'):
                has_docx = True
            elif item.endswith('.mxliff'):
                has_mxliff = True
        if has_docx and has_mxliff:
            break
    if has_docx and has_mxliff:
        print("\nCheck 1: The directory contains at least one .docx file and one .mxliff file.")
    else:
        print("Check 1:")
        if not has_docx:
            print("No .docx file found. Add file(s) to MagicBox.")
        if not has_mxliff:
            print("No .mxliff file found. Add file(s) to MagicBox.")

# Process MXLIFF files
def parse_mxliff_to_df(mxliff_file):
    print("Processing .MXLIFF file...")
    ET.register_namespace('m', 'urn:oasis:names:tc:xliff:document:1.2')
    tree = ET.parse(mxliff_file)
    root = tree.getroot()
    namespaces = {'m': 'urn:oasis:names:tc:xliff:document:1.2'}

    sources = []
    targets = []
    match_qualities = []

    for trans_unit in root.findall('.//m:trans-unit', namespaces):
        source_text = trans_unit.find('m:source', namespaces).text if trans_unit.find('m:source', namespaces) is not None else ''
        target_text = trans_unit.find('m:target', namespaces).text if trans_unit.find('m:target', namespaces) is not None else ''
        match_quality = '0'

        for alt_trans in trans_unit.findall('.//m:alt-trans', namespaces):
            if alt_trans.attrib.get('origin') == 'memsource-tm':
                match_quality = int(float(alt_trans.attrib.get('match-quality', '0')) * 100)
                break

        sources.append(source_text)
        targets.append(target_text)
        match_qualities.append(match_quality)

    df = pd.DataFrame({
        'Source': sources,
        'Target': targets,
        'Match': match_qualities
    })

    df.index += 1
    df.reset_index(inplace=True)
    df.rename(columns={'index': 'Index'}, inplace=True)
    df['Comment'] = ""
    return df

# Function to format tables
def format_table(table):
    t_settings = CONFIG["TableFormattingSettings"]
    table.style = 'Table Grid'
    row_widths = t_settings["RowWidths"]

    for i, width in enumerate(row_widths):
        for cell in table.columns[i].cells:
            cell.width = Mm(width)

    blue_color = "95B3D7"
    first_column_cells = table.rows[0].cells
    change_cell_color(first_column_cells, blue_color)

# Helper functions
def change_cell_color(cells, background_color):
    for cell in cells:
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), background_color)
        tcPr.append(shd)

# Function to process Word files
def process_word_file(file_path, output_folder):
    p_settings = CONFIG["ProcessingSettings"]
    doc = Document(file_path)

    tables_to_delete = p_settings["DeleteFirstNTables"]
    for _ in range(tables_to_delete):
        if len(doc.tables) > 0:
            table = doc.tables[0]
            table._element.getparent().remove(table._element)

    final_col_length = len(CONFIG["GeneralSettings"]["Column_Headers"])
    original_table = doc.tables[0]
    new_table = doc.add_table(rows=0, cols=final_col_length)

    # Copy content and process...
    df_table = table_to_df(new_table)
    return df_table

def table_to_df(table):
    column_headers = CONFIG["GeneralSettings"]["Column_Headers"]
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        data.append(row_data)
    return pd.DataFrame(data, columns=column_headers)

# Main execution block
def main():
    magic_box_path = '/content/drive/MyDrive/MagicBox/'
    output_folder_path = os.path.join(magic_box_path, 'Output_Folder')
    config_folder_path = os.path.join(magic_box_path, 'configFiles')
    config_file_path = os.path.join(config_folder_path, 'config.json')

    ensure_folder_exists(output_folder_path)
    ensure_folder_exists(config_folder_path)

    check_for_files(magic_box_path)

    if os.path.exists(config_file_path):
        print("\nCheck 3: Configuration file OK\n")
    else:
        print("\nCheck 3: Configuration file missing from MagicBox/configFolder!")

if __name__ == "__main__":
    main()
