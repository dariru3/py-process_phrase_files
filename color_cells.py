from docx import Document
import os
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def change_table_cell(cell, background_color=None):
    """Changes the background color of a table cell."""
    if background_color:
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
        cell._tc.get_or_add_tcPr().append(shading_elm)

def process_word_file(file_path):
    doc = Document(file_path)

    # Delete the first three tables
    for _ in range(3):
        if len(doc.tables) > 0:
            tbl = doc.tables[0]
            tbl._element.getparent().remove(tbl._element)

    # Process the fourth table (now the first table in the document)
    if len(doc.tables) > 0:
        original_table = doc.tables[0]

        # Create a new table with 4 columns
        new_table = doc.add_table(rows=0, cols=4)

        for row in original_table.rows:
            new_row = new_table.add_row()
            new_cells = new_row.cells
            # Copy the content from the original table
            new_cells[0].text = row.cells[2].text  # 3rd column: numbers
            new_cells[1].text = row.cells[3].text  # 4th column: japanese
            new_cells[2].text = row.cells[5].text  # 5th column: english
            new_cells[3].text = row.cells[6].text  # 6th column: matching percent

            # If the english column cell has text, highlight cells gray
            if new_cells[2].text.strip():
                change_table_cell(new_cells[1], background_color="D9D9D9")  # Gray color
                change_table_cell(new_cells[2], background_color="D9D9D9")

        # Remove the original table
        original_table._element.getparent().remove(original_table._element)

    # Save the document
    new_file_path = file_path.replace('.docx', '_processed.docx')
    doc.save(new_file_path)

    print(f"Processed file saved as: {new_file_path}")

def process_all_word_files_in_folder(folder_path):
    for file_name in os.listdir(folder_path):
        if file_name.endswith('.docx'):
            file_path = os.path.join(folder_path, file_name)
            process_word_file(file_path)

# Example usage
input_folder = 'input_files'
process_all_word_files_in_folder(input_folder)
