import os
import unittest
from docx import Document
from src.save_formatting import extract_formatting_from_column
from src.process_mxliff import remove_tags

class TestDocxMerge(unittest.TestCase):
    input_folder = "data/input_files"
    input_files = [f for f in os.listdir(input_folder) if f.endswith(".docx")]
    INPUT_PATH = os.path.join(input_folder, input_files[0])
    OUTPUT_PATH = os.path.join(
        "data/output_files",
        f"{os.path.splitext(os.path.basename(INPUT_PATH))[0]}_merged.docx"
    )

    # 3rd table in input, 1st table in output
    INPUT_TABLE_INDEX = 3
    OUTPUT_TABLE_INDEX = 0

    # map input‐table cols → output‐table cols
    COL_MAP = {
        3: 1,  # input col 3 → output col 1 (Source/Japanese)
        5: 2,  # input col 5 → output col 2 (Target/English)
        7: 3   # input col 7 → output col 4 (Comment)
    }

    def setUp(self):
        self.input_doc = Document(self.INPUT_PATH)
        self.output_doc = Document(self.OUTPUT_PATH)
        self.in_table  = self.input_doc.tables[self.INPUT_TABLE_INDEX]
        self.out_table = self.output_doc.tables[self.OUTPUT_TABLE_INDEX]
        self.col_map = TestDocxMerge.COL_MAP

    def test_row_count(self):
        self.assertEqual(
            len(self.in_table.rows) + 1, # Input table does not have the header row
            len(self.out_table.rows),
            f"Row count mismatch: input has {len(self.in_table.rows)}, "
            f"output has {len(self.out_table.rows)}"
        )

    def test_cell_texts_mirror(self):
        for row_idx, (in_row, out_row) in enumerate(
            zip(self.in_table.rows, self.out_table.rows[1:])
        ):
            for in_col, out_col in self.col_map.items():
                in_text  = in_row.cells[in_col].text.strip()
                out_text = out_row.cells[out_col].text.strip()
                self.assertEqual(
                    remove_tags(in_text), remove_tags(out_text),
                    f"Text mismatch at row {row_idx + 1}, input col {in_col}, output col {out_col + 1}"
                )

    def test_formatting_info_mirror(self):
        keys_to_check = [c for c in self.col_map.keys() if c != 7]
        vals_to_check = [self.col_map[c] for c in keys_to_check]

        in_fmt = extract_formatting_from_column(
            self.input_doc,  self.INPUT_TABLE_INDEX, keys_to_check
        )
        out_fmt = extract_formatting_from_column(
            self.output_doc, self.OUTPUT_TABLE_INDEX, vals_to_check
        )

        for row_idx in in_fmt:
            for in_col in keys_to_check:
                out_col = self.col_map[in_col]
                runs_in  = in_fmt[row_idx][in_col]
                runs_out = out_fmt[row_idx + 1][out_col]

                # Drop blank text runs
                def clean_runs(runs):
                    processed_runs = [
                        {**run, "text": remove_tags(run["text"])} for run in runs
                    ]
                    return [run for run in processed_runs if run["text"].strip()]

                clean_in = clean_runs(runs_in)
                clean_out = clean_runs(runs_out)

                self.assertEqual(
                    clean_in, clean_out,
                    f"Formatting mismatch at row {row_idx + 1}, "
                    f"input col {in_col + 1}, output col {out_col + 1}"
                )

# Use `python3 -m unittest tests/test_main` to run file from console
if __name__ == "__main__":
    unittest.main()
