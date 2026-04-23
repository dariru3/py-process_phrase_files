import contextlib
import io
import os
import subprocess
import sys
import tempfile
import unittest

from docx import Document
from docx.shared import Pt, RGBColor

from src.df_to_word import get_file_pairs
from src.pipeline import run_pipeline
from src.process_mxliff import parse_mxliff_to_df
from src.save_formatting import extract_formatting_from_column


class TestPipeline(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.input_dir = os.path.join(self.temp_dir.name, "input")
        self.output_dir = os.path.join(self.temp_dir.name, "output")
        os.makedirs(self.input_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_parse_mxliff_handles_empty_source_tag_cleanup_and_match_quality(self):
        mxliff_path = os.path.join(self.input_dir, "parser_case.mxliff")
        self._write_mxliff(
            mxliff_path,
            """<?xml version="1.0" encoding="UTF-8"?>
<xliff xmlns="urn:oasis:names:tc:xliff:document:1.2" version="1.2">
  <file>
    <body>
      <trans-unit id="1">
        <source>Alpha {^&gt;2&lt;^} source{j}</source>
        <target>Target line one</target>
        <alt-trans origin="memsource-tm" match-quality="1.0" />
      </trans-unit>
      <trans-unit id="2">
        <source />
        <target>Target line two</target>
      </trans-unit>
    </body>
  </file>
</xliff>
""",
        )

        df = parse_mxliff_to_df(mxliff_path)

        self.assertEqual(["ID", "Source", "Target", "Match", "Comment"], list(df.columns))
        self.assertEqual("1", df.loc[0, "ID"])
        self.assertEqual("Alpha 2 source", df.loc[0, "Source"])
        self.assertEqual("Target line one", df.loc[0, "Target"])
        self.assertEqual(100, df.loc[0, "Match"])
        self.assertEqual("", df.loc[0, "Comment"])
        self.assertEqual("2", df.loc[1, "ID"])
        self.assertEqual("", df.loc[1, "Source"])
        self.assertEqual("Target line two", df.loc[1, "Target"])
        self.assertEqual(0, df.loc[1, "Match"])

    def test_generated_fixtures_cover_pairing_pipeline_output_and_formatting(self):
        base_name = "generic_case"
        docx_path = os.path.join(self.input_dir, f"{base_name}.docx")
        mxliff_path = os.path.join(self.input_dir, f"{base_name}.mxliff")
        self._build_input_docx(docx_path)
        self._write_mxliff(
            mxliff_path,
            """<?xml version="1.0" encoding="UTF-8"?>
<xliff xmlns="urn:oasis:names:tc:xliff:document:1.2" version="1.2">
  <file>
    <body>
      <trans-unit id="1">
        <source>Alpha one</source>
        <target>Target one</target>
        <alt-trans origin="memsource-tm" match-quality="1.0" />
      </trans-unit>
      <trans-unit id="2">
        <source>Gamma 2</source>
        <target>Fallback target two</target>
      </trans-unit>
    </body>
  </file>
</xliff>
""",
        )
        self._build_extra_docx(os.path.join(self.input_dir, "orphan.docx"))
        self._write_mxliff(
            os.path.join(self.input_dir, "lonely.mxliff"),
            """<?xml version="1.0" encoding="UTF-8"?>
<xliff xmlns="urn:oasis:names:tc:xliff:document:1.2" version="1.2">
  <file><body><trans-unit id="9"><source>Unused</source><target>Unused</target></trans-unit></body></file>
</xliff>
""",
        )

        pairs = get_file_pairs(self.input_dir)
        self.assertEqual([(f"{base_name}.docx", f"{base_name}.mxliff")], pairs)

        processed_count = self._run_pipeline(self.input_dir, self.output_dir)
        self.assertEqual(1, processed_count)

        self.assertEqual(0, self._run_pipeline(self.input_dir, self.output_dir))
        self.assertEqual(
            1,
            self._run_pipeline(self.input_dir, self.output_dir, skip_existing=False),
        )

        output_path = os.path.join(self.output_dir, f"{base_name}_merged.docx")
        self.assertTrue(os.path.exists(output_path))

        input_doc = Document(docx_path)
        output_doc = Document(output_path)
        input_table = input_doc.tables[3]
        output_table = output_doc.tables[0]

        self.assertEqual(
            ["p", "Japanese", "English", "Comment"],
            [cell.text for cell in output_table.rows[0].cells],
        )
        self.assertEqual(len(input_table.rows) + 1, len(output_table.rows))

        output_rows = output_table.rows[1:]
        self.assertEqual("1", output_rows[0].cells[0].text.strip())
        self.assertEqual("Alpha one", output_rows[0].cells[1].text.strip())
        self.assertEqual("Target one", output_rows[0].cells[2].text.strip())
        self.assertEqual("note a", output_rows[0].cells[3].text.strip())
        self.assertEqual("2", output_rows[1].cells[0].text.strip())
        self.assertEqual("Gamma 2", output_rows[1].cells[1].text.strip())
        self.assertEqual("Fallback target two", output_rows[1].cells[2].text.strip())
        self.assertEqual("note b", output_rows[1].cells[3].text.strip())

        input_formatting = extract_formatting_from_column(input_doc, 3, [3, 5])
        output_formatting = extract_formatting_from_column(output_doc, 0, [1, 2])

        self.assertEqual(
            self._clean_runs(input_formatting[0][3]),
            self._clean_runs(output_formatting[1][1]),
        )
        self.assertEqual(
            self._clean_runs(input_formatting[0][5]),
            self._clean_runs(output_formatting[1][2]),
        )
        self.assertEqual(
            self._clean_runs(input_formatting[1][3]),
            self._clean_runs(output_formatting[2][1]),
        )

    def test_cli_help_returns_promptly(self):
        result = subprocess.run(
            [sys.executable, "main.py", "--help"],
            cwd=os.getcwd(),
            capture_output=True,
            text=True,
            timeout=5,
            check=True,
        )

        self.assertIn("usage: main.py", result.stdout)
        self.assertIn("--input", result.stdout)
        self.assertIn("--output", result.stdout)

    def _build_input_docx(self, path):
        document = Document()

        for index in range(3):
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = f"placeholder {index + 1}"

        table = document.add_table(rows=0, cols=8)
        row_one = table.add_row().cells
        row_one[0].text = "1"
        self._add_runs(
            row_one[3],
            [
                {"text": "Alpha ", "bold": True, "font_name": "Calibri", "font_size": 13, "font_color": "AA0000"},
                {"text": "one", "italic": True, "underline": True, "font_name": "Calibri", "font_size": 11},
            ],
        )
        row_one[5].text = ""
        self._add_runs(
            row_one[5],
            [
                {"text": "Target ", "underline": True, "font_name": "Times New Roman", "font_size": 12},
                {"text": "one", "superscript": True, "font_name": "Times New Roman", "font_size": 9, "font_color": "0000AA"},
            ],
        )
        row_one[6].text = "0"
        row_one[7].text = "note a"

        row_two = table.add_row().cells
        row_two[0].text = "2"
        self._add_runs(
            row_two[3],
            [
                {"text": "Gamma ", "font_name": "Arial", "font_size": 12},
                {"text": "2", "subscript": True, "font_name": "Arial", "font_size": 8},
            ],
        )
        row_two[5].text = ""
        row_two[6].text = ""
        row_two[7].text = "note b"

        document.save(path)

    def _build_extra_docx(self, path):
        document = Document()
        document.add_paragraph("unmatched file")
        document.save(path)

    def _add_runs(self, cell, run_specs):
        paragraph = cell.paragraphs[0]
        for spec in run_specs:
            run = paragraph.add_run(spec["text"])
            run.bold = spec.get("bold")
            run.italic = spec.get("italic")
            run.underline = spec.get("underline")
            run.font.name = spec.get("font_name")
            if spec.get("font_size") is not None:
                run.font.size = Pt(spec["font_size"])
            if spec.get("font_color"):
                run.font.color.rgb = RGBColor.from_string(spec["font_color"])
            run.font.superscript = spec.get("superscript")
            run.font.subscript = spec.get("subscript")

    def _write_mxliff(self, path, content):
        with open(path, "w", encoding="utf-8") as handle:
            handle.write(content)

    def _run_pipeline(self, input_dir, output_dir, skip_existing=True):
        with contextlib.redirect_stdout(io.StringIO()):
            return run_pipeline(input_dir, output_dir, skip_existing=skip_existing)

    def _clean_runs(self, runs):
        cleaned = []
        for run in runs:
            if not run["text"].strip():
                continue
            cleaned.append(
                {
                    "text": run["text"],
                    "bold": run["bold"],
                    "italic": run["italic"],
                    "underline": run["underline"],
                    "font_name": run["font_name"],
                    "font_size": run["font_size"],
                    "font_color": run["font_color"],
                    "superscript": run["superscript"],
                    "subscript": run["subscript"],
                }
            )
        return cleaned


if __name__ == "__main__":
    unittest.main()
