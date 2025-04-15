# @title Step 2: Run Magic Box
from datetime import datetime
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.shared import Pt, RGBColor
import os
import pandas as pd
import re
import xml.etree.ElementTree as ET

# Start of src/config_loader.py
CONFIG = {
    "GeneralSettings": { # When updating Colab, replace with commented folder paths
        "InputFolderPath": "data/input_files/", # "/content/drive/MyDrive/MagicBox/",
        "OutputFolderPath": "data/output_files/", # "/content/drive/MyDrive/MagicBox/Output_Folder/",
        "Column_Headers": ["ID", "Source", "Target", "Match", "Comment"]
    },
    "ProcessingDocSettings": {
        "DeleteFirstNTables": 3,
        "ColumnsToKeep": [0, 3, 5, 6, 7]
    },
    "ProcessingXliffSettings": {
        "TagPatterns": r"\{.?>|<.?\}|\{j\}", # Remove custom tags such as {b>, <b}, {j} from the input text.
        "XliffNamespace": "urn:oasis:names:tc:xliff:document:1.2",
    },
    "ConditionalFormattingSettings": {
        "TargetColumnIndex": 2,
        "MatchColumnIndex": 3,
        "CommentColumnIndex": 4,
        "CommentToGray": ["lock", "locked"],
        "MatchToGray": ["100", "101"],
        "BackgroundColor": "D9D9D9"
    },
    "TableFormattingSettings": {
        "RowWidths": [9, 81, 112, 11, 21],
        "NewColumnNames": {'ID': 'p', 'Source': 'Japanese', 'Target': 'English'}
    }
}

# End of src/config_loader.py


# Start of src/save_formatting.py

def extract_formatting_from_column(doc, table_num, col_nums):
    table = doc.tables[table_num]
    formatting_info = {}

    for row_idx, row in enumerate(table.rows):
        row_format = {}
        for col in col_nums:
            cell = row.cells[col]
            cell_info = []
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run_info = {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name,
                        "font_size": run.font.size.pt if run.font.size else None,
                        "font_color": run.font.color.rgb if run.font.color else None,
                        "superscript": run.font.superscript,
                        "subscript": run.font.subscript,
                    }
                    cell_info.append(run_info)
            row_format[col] = cell_info
        formatting_info[row_idx] = row_format

    return formatting_info

def reapply_formatting_to_column(table, formatting_info, col_nums, table_num=0):
    col_mapping = { 3: 1, 5: 2 } # Maps original column indices (3 and 5) to new column indices (1 and 2)
    for row_idx, cols_info in formatting_info.items():
        for orig_col, new_col in col_mapping.items():
            cell_info = cols_info.get(orig_col, [])
            has_previous_text = any(run_info["text"] for run_info in cell_info)
            if not has_previous_text:
                continue

            cell = table.cell(row_idx + 1, new_col) # Skip header row
            cell.text = ""
            for run_info in cell_info:
                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                run = paragraph.add_run(run_info["text"])
                run.bold = run_info.get("bold")
                run.italic = run_info.get("italic")
                run.underline = run_info.get("underline")
                run.font.name = run_info.get("font_name")
                if run_info.get("font_size"):
                    run.font.size = Pt(run_info["font_size"])
                if run_info.get("font_color"):
                    run.font.color.rgb = RGBColor.from_string(str(run_info["font_color"]))
                else:
                    run.font.color.rgb = None
                run.font.superscript = run_info.get("superscript")
                run.font.subscript = run_info.get("subscript")

# End of src/save_formatting.py


# Start of src/format_helper.py

def apply_superscript(run, text):
    """Helper function to set text as superscript"""
    run.text = text
    run.font.superscript = True

def format_superscripts(paragraph):
    """Helper function to format text with superscript tags"""
    text = paragraph.text
    parts = re.split(r'(\{\^\>.*?\<\^\}|\{.*?\>.*?\<.*?\})', text)

    new_runs = []
    for part in parts:
        if part.startswith('{^>') and part.endswith('<^}'):
            # This is a superscript
            superscript_text = part[3:-3]
            run = paragraph.add_run()
            apply_superscript(run, superscript_text)
            new_runs.append(run)
        elif re.match(r'\{.*?\>.*?\<.*?\}', part):
            # Handle other custom tags if necessary
            inner_text = part.split('>')[1].split('<')[0]
            run = paragraph.add_run()
            apply_superscript(run, inner_text)
            new_runs.append(run)
        else:
            # Normal text
            run = paragraph.add_run(part)
            new_runs.append(run)

    # Remove original runs
    for run in paragraph.runs:
        p = run._element.getparent()
        p.remove(run._element)

    # Append new runs with formatting
    for new_run in new_runs:
        run = paragraph.add_run(new_run.text)
        run.font.superscript = new_run.font.superscript

def reformat_text(table):
    """Process a single table in the document"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                format_superscripts(paragraph)

def change_cell_color(cells, background_color):
    for cell in cells:
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), background_color)
        tcPr.append(shd)

def set_column_language(table, column_index, language_code):
    ''''
    Not working as intended
    '''
    for row in table.rows:
        cell = row.cells[column_index]
        for paragraph in cell.paragraphs:
            rPr = paragraph.runs[0].element.get_or_add_rPr()
            lang = OxmlElement('w:lang')
            lang.set(qn('w:val'), language_code)
            rPr.append(lang)

# format column widths and header row cell color
def format_table(table):
    t_settings = CONFIG["TableFormattingSettings"]
    table.style = 'Table Grid'
    row_widths = t_settings["RowWidths"]

    for i, width in enumerate(row_widths):
        for cell in table.columns[i].cells:
            cell.width = Mm(width)

    blue_color = "95B3D7"  # Hex code for blue
    first_column_cells = table.rows[0].cells
    change_cell_color(first_column_cells, blue_color)

def apply_conditional_formatting(table):
    '''
    Change row cell colors to gray if either condition is met.
    Condition 1: There is text in the target cell and the match is either "100" or "101"
    Condition 2: There is text in the target cell and the comment is "lock" or "locked"
    '''
    c_settings = CONFIG["ConditionalFormattingSettings"]
    target_column_index = c_settings["TargetColumnIndex"] # 2
    match_column_index = c_settings["MatchColumnIndex"] #3
    comment_column_index = c_settings["CommentColumnIndex"] # 4
    comment_to_gray = c_settings["CommentToGray"] # ['lock', 'locked']
    match_to_gray = c_settings["MatchToGray"] # ['100', '101']
    background_color= c_settings["BackgroundColor"] # "D9D9D9"

    for row in table.rows:
        target_value = row.cells[target_column_index].text.strip()
        match_value = row.cells[match_column_index].text.strip()
        comment_value = row.cells[comment_column_index].text.lower().strip()

        condition_1_met = target_value and match_value in match_to_gray
        condition_2_met = target_value and comment_value in comment_to_gray

        if condition_1_met or condition_2_met:
            cells_to_color = [cell for cell in row.cells]
            change_cell_color(cells_to_color, background_color)

def set_landscape_orientation(document):
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297) # A4 width
    section.page_height = Mm(210) # A4 height

# set font size and line spacing
def format_font_lines(document):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    line_space = 1.15
    column_1_font_size = 8

    for paragraph in document.paragraphs:
        apply_paragraph_format(paragraph, style, line_space)

    for table in document.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if i == 0: # set font size of column 1 to 8pt
                        apply_paragraph_format(paragraph, style, line_space, column_1_font_size)
                    else:
                        apply_paragraph_format(paragraph, style, line_space)

# helper function for format_font_lines()
def apply_paragraph_format(paragraph, style, line_space, font_size=None):
    paragraph.style = style
    paragraph.paragraph_format.line_spacing = line_space
    if font_size:
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(font_size)

def apply_formatting_pipe(table, doc):
    format_table(table)
    apply_conditional_formatting(table)
    set_column_language(table, 1, "ja-JP")
    set_landscape_orientation(doc)
    format_font_lines(doc)

# End of src/format_helper.py


# Start of src/process_word.py

def delete_first_n_tables(doc, n):
    for _ in range(n):
        if len(doc.tables) > 0:
            table = doc.tables[0]
            table._element.getparent().remove(table._element)

def copy_content_to_table(original_table, new_table, columns_to_copy):
    for row in original_table.rows:
        new_row = new_table.add_row()
        new_cells = new_row.cells
        for i, col_index in enumerate(columns_to_copy):
            original_text = row.cells[col_index].text
            cleansed_text = remove_tags(original_text)
            new_cells[i].text = cleansed_text

def process_word_file(file_path, output_folder):
    # Load settings
    p_settings = CONFIG["ProcessingDocSettings"]
    final_col_length = len(CONFIG["GeneralSettings"]["Column_Headers"])

    print("Processing .DOCX file...")
    doc = Document(file_path)

    # Save text formatting to reapply later
    formatting_info = extract_formatting_from_column(doc, 3, [3, 5])

    # Remove other tables
    tables_to_delete = p_settings["DeleteFirstNTables"]
    delete_first_n_tables(doc, tables_to_delete)

    # Get contents from source table from certain columns
    original_table = doc.tables[0]
    new_table = doc.add_table(rows=0, cols=final_col_length)
    columns_to_copy = p_settings["ColumnsToKeep"]
    copy_content_to_table(original_table, new_table, columns_to_copy)

    df_table = table_to_df(new_table)
    return df_table, formatting_info

# End of src/process_word.py


# Start of src/process_mxliff.py

def remove_tags(text):
    patterns = CONFIG["ProcessingXliffSettings"]["TagPatterns"]
    cleansed_text = re.sub(patterns, '', text)

    return cleansed_text

def setup_root(mxliff_file, xliff_namespace):
    """
    Register the given namespace and parse the MXLIFF file,
    returning the root element.
    """
    ET.register_namespace('m', xliff_namespace)
    tree = ET.parse(mxliff_file)
    return tree.getroot()

def get_match_quality(alt_trans):
    """
    Extract and calculate the match quality from an alt-trans element.
    Returns the match quality as an integer percentage.
    """
    if alt_trans.attrib.get('origin') == 'memsource-tm':
        match_quality = alt_trans.attrib.get('match-quality', '0')
        return int(float(match_quality) * 100)
    else:
        return 0

def parse_mxliff_to_df(mxliff_file):
    print("Processing .MXLIFF file...")
    xliff_namespace = CONFIG["ProcessingXliffSettings"]["XliffNamespace"]
    root = setup_root(mxliff_file, xliff_namespace)

    # Define the namespaces used in your MXLIFF file
    namespaces = {'m': xliff_namespace}

    # Initialize lists to hold the extracted data
    ids = []
    sources = []
    targets = []
    match_qualities = []

    # Loop through each translation unit in the MXLIFF file
    find_text = lambda trans_unit, m : trans_unit.find(m, namespaces).text if trans_unit.find(m, namespaces) is not None else ''

    for trans_unit in root.findall('.//m:trans-unit', namespaces):
        unit_id = trans_unit.attrib.get('id', '')
        source_text = find_text(trans_unit, 'm:source')
        source_text = remove_tags(source_text)
        target_text = find_text(trans_unit, 'm:target')

        # Check for alt-trans elements with origin="memsource-tm" and extract match-quality
        match_quality = 0 # Default value
        for alt_trans in trans_unit.findall('.//m:alt-trans', namespaces):
            match_quality = get_match_quality(alt_trans)
            if match_quality != 0:
                break  # Assuming we only need the first matching alt-trans entry

        ids.append(unit_id)
        sources.append(source_text)
        targets.append(target_text)
        match_qualities.append(match_quality)

    # Create a DataFrame from the extracted data
    df = pd.DataFrame({
        'ID': ids,
        'Source': sources,
        'Target': targets,
        'Match': match_qualities
    })

    df['Comment'] = ""

    return df

# End of src/process_mxliff.py


# Start of src/table_to_df.py

def table_to_df(table):
    column_headers = CONFIG["GeneralSettings"]["Column_Headers"]
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        data.append(row_data)
    return pd.DataFrame(data, columns=column_headers)

# End of src/table_to_df.py


# Start of src/merge_df.py

def merge_dfs(df_word, df_mxliff):
    # Convert Match to int, with error handling
    df_word['Match'] = pd.to_numeric(df_word['Match'], errors='coerce').fillna(0).astype(int)
    df_mxliff['Match'] = pd.to_numeric(df_mxliff['Match'], errors='coerce').fillna(0).astype(int)

    # Merge the DataFrames
    df_combined = pd.merge(df_word, df_mxliff, on=['ID'], how='outer', suffixes=('', '_df2'))

    # Select the best values for each column based on availability and preference
    df_combined['Target'] = df_combined['Target'].where(df_combined['Target'] != '', df_combined['Target_df2'])

    df_combined['Match'] = df_combined['Match'].fillna(0).astype(int)
    df_combined['Match'] = df_combined['Match'].where(df_combined['Match'] != 0, df_combined['Match_df2']).fillna(0).astype(int)

    df_combined['Comment'] = df_combined['Comment'].where(df_combined['Comment'] != '', df_combined['Comment_df2'])

    # Drop the temporary columns from df2
    df_combined.drop(columns=['Source_df2', 'Target_df2', 'Match_df2', 'Comment_df2'], inplace=True)

    return df_combined

# End of src/merge_df.py


# Start of src/df_to_word.py

def delete_column_in_table(table, column_index=CONFIG["ConditionalFormattingSettings"]["MatchColumnIndex"]):
    grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
    for cell in table.column_cells(column_index):
        cell._tc.getparent().remove(cell._tc)
    col_elem = grid[column_index]
    grid.remove(col_elem)

def get_file_pairs(folder_path):
    docx_files = {}
    mxliff_files = {}
    for filename in os.listdir(folder_path):
        base_name, ext = os.path.splitext(filename)
        if ext == ".docx":
            docx_files[base_name] = filename
        elif ext == ".mxliff":
            mxliff_files[base_name] = filename

    # Pairing files with the same base name
    pairs = []
    for base_name, docx_file in docx_files.items():
        mxliff_file = mxliff_files.get(base_name)
        if mxliff_file:
            pairs.append((docx_file, mxliff_file))
    return pairs

def setup_table(df, table):
    t_settings = CONFIG["TableFormattingSettings"]
    line_num_col = "p"

    # Rename column headers
    df.rename(columns=t_settings["NewColumnNames"], inplace=True)
    # Reassign the 'p' column from id numbers to index numbers
    df[line_num_col] = range(1, len(df) + 1)

    # Add header row to new table
    for i, column in enumerate(df.columns):
        table.cell(0, i).text = str(column)

    # Add data rows to new table
    for i, row in df.iterrows():
        cells = table.add_row().cells
        for j, value in enumerate(row):
            if pd.isnull(value) or value == "None":
                cells[j].text = ""
            else:
                cells[j].text = str(value)

def save_doc_file(docx_file, new_doc, output_folder):
    # Check output folder
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Use original .docx file name, add "_merged"
    output_file_path = os.path.join(output_folder, f"{os.path.splitext(docx_file)[0]}_merged.docx")

    # Add metadata
    new_doc.core_properties.created = datetime.now()
    new_doc.core_properties.modified = datetime.now()

    # Save file and notify in console
    new_doc.save(output_file_path)
    print(f"Merged tables saved as Word document: {output_file_path}.")

def dataframe_to_word_table(docx_file, df, output_folder, formatting_info):
    # Create new .docx file with a new blank table
    new_doc = Document()
    table = new_doc.add_table(rows=1, cols=len(df.columns))
    table.autofit = False

    setup_table(df, table)
    apply_formatting_pipe(table, new_doc) # Format table

    # Reapply formatting to Japanese and English text
    reapply_formatting_to_column(table, formatting_info, [3, 5])

    # Drop the 'Match' column after all formatting is done
    delete_column_in_table(table)

    # Save new .docx file
    save_doc_file(docx_file, new_doc, output_folder)

def process_files(docx_file, mxliff_file, input_folder, output_folder):
    df_word, formatting_info = None, None

    # Process the Word files
    df_word, formatting_info = process_word_file(os.path.join(input_folder, docx_file), output_folder)

    # Check Word data, then process MXLIFF files
    if df_word is not None and not df_word.empty and formatting_info is not None:
        df_mxliff = parse_mxliff_to_df(os.path.join(input_folder, mxliff_file))
    else:
        print("Failed to process Word file.")
        return

    # Merge the DataFrames
    merged_df = merge_dfs(df_word, df_mxliff)

    # Save the merged DataFrame to a Word document
    dataframe_to_word_table(docx_file, merged_df, output_folder, formatting_info)

# End of src/df_to_word.py


# Start of scripts/main.py

def filter_unprocessed_pairs(pairs, output_folder):
    unprocessed_pairs = []
    for docx_file, mxliff_file in pairs:
        # Extract the base name from the docx file (base name for mxliff file should be the same)
        base_name, _ = os.path.splitext(os.path.basename(docx_file))

        # Check if the merged file exists
        merged_filename = f"{base_name}_merged.docx"
        merged_file_path = os.path.join(output_folder, merged_filename)
        if not os.path.exists(merged_file_path):
            unprocessed_pairs.append((docx_file, mxliff_file))
        else:
            print(f"Skipped processing for {base_name} because the merged file already exists.")
    return unprocessed_pairs

def main():
    g_settings = CONFIG["GeneralSettings"]
    input_folder = g_settings["InputFolderPath"] # "input_files/"
    output_folder = g_settings["OutputFolderPath"] # "output_files/"

    pairs = get_file_pairs(input_folder)
    unprocessed_pairs = filter_unprocessed_pairs(pairs, output_folder)
    for docx_file, mxliff_file in unprocessed_pairs:
        print(f"File pair:\n{docx_file}\n{mxliff_file}")
        process_files(docx_file, mxliff_file,input_folder, output_folder)

# Use `python3 -m scripts.main` to run file from console
if __name__ == "__main__":
    main()

# End of scripts/main.py

