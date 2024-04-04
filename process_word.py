from docx import Document
from table_to_df import table_to_df
import os
import re
from config_loader import CONFIG

def delete_first_n_tables(doc, n):
    for _ in range(n):
        if len(doc.tables) > 0:
            table = doc.tables[0]
            table._element.getparent().remove(table._element)

def copy_content_to_table(original_table, new_table, columns_to_copy):
    for row in original_table.rows:
        new_row = new_table.add_row()
        new_cells = new_row.cells
        for i, col_index in enumerate(columns_to_copy):
            new_cells[i].text = row.cells[col_index].text

def process_word_file(file_path, output_folder, attempts=1):
    p_settings = CONFIG["ProcessingSettings"]
    final_col_length = len(CONFIG["GeneralSettings"]["Column_Headers"])
    if attempts == 1:
        print("Processing .DOCX file...")
    max_attempts = p_settings["MaxAttempts"]
    doc = Document(file_path)

    tables_to_delete = p_settings["DeleteFirstNTables"]
    delete_first_n_tables(doc=doc, n=tables_to_delete)

    columns_to_copy = adjust_columns_by_attempts(attempts, p_settings)

    original_table = doc.tables[0]
    new_table = doc.add_table(rows=0, cols=final_col_length)

    copy_content_to_table(original_table, new_table, columns_to_copy)

    if validate_table_contents(new_table, p_settings):
        df_table = table_to_df(new_table)
        print("Success!")
        return df_table
    else:
        if attempts < max_attempts:
            print(f'Attempt {attempts} failed, trying again...')
            return process_word_file(file_path, output_folder, attempts + 1)
        else:
            print(f'Maximum attempts reached for file {file_path}. File processing aborted.')
            return None

def contains_japanese(text, process_settings):
    # Regular expression for matching Japanese characters
    pattern = process_settings["JapanesePattern"]
    return re.search(pattern, text) is not None

def validate_table_contents(new_table, process_settings):
    valid_rows = True
    for i, row in enumerate(new_table.rows[1:11]):
        column_3_target_text = row.cells[2].text

        if column_3_target_text and contains_japanese(column_3_target_text, process_settings):
            print(f"Invalid row {i}: {column_3_target_text}")
            valid_rows = False
    
    return valid_rows

def adjust_columns_by_attempts(attempts, process_settings):
    attempt_1_col = process_settings["Mapping_1"]
    attempt_2_col = process_settings["Mapping_2"]
    attempts_mapping = {
        1: ("First attempt", attempt_1_col),
        2: ("Second attempt", attempt_2_col),
    }

    message, columns = attempts_mapping.get(attempts, ("Second attempt failed", None))
    print(message)
    return columns